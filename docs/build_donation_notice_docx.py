# -*- coding: utf-8 -*-
"""Generate docs/大庙捐款支持公示.docx（收支表 + 支出明细表）.

Run:
  python docs/build_donation_notice_docx.py
  python docs/build_donation_notice_docx.py --sync-md-table   # 按 ROWS 重写 MD 中「支出明细表」HTML（类别 rowspan）
  python docs/build_donation_notice_docx.py --print-expense-html

Word 中方框「□」多为未指定东亚字体（eastAsia）：仅设置 run.font.name 往往只作用于西文，
中文会落到无字形字体上。生成后统一为每个 run 写入 w:eastAsia（见 fix_all_runs_cjk_font）。
"""

from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

OUT_DIR = Path(__file__).resolve().parent

# 简体中文 Windows / Office 普遍自带；若打开仍异常可改为「宋体」
CJK_FONT = "微软雅黑"

ROWS = [
    ("神台与塑像", "神台建造", "", "13,000.00"),
    ("神台与塑像", "五尊塑像费用及装藏用品", "", "48,162.00"),
    ("神台与塑像", "装藏辅料（棉、草、铜钱、花、药、线、镜、铃等）", "出于敬畏不逐项展开，作一笔良心账", "2,982.00"),
    ("开光用品", "脸盆、镜子、毛巾、梳子", "", "87.50"),
    ("开光用品", "大红水果供盘", "8 元 × 5 个", "40.00"),
    ("开光用品", "白瓷碗", "4 元 × 2 个", "8.00"),
    ("开光用品", "红布", "16 尺", "16.00"),
    ("开光用品", "黄布", "16 尺", "16.00"),
    ("开光用品", "好无烟香", "15 元 × 6 把", "90.00"),
    ("开光用品", "罗汉裱", "2 元 × 6 把（实收 10 元）", "10.00"),
    ("开光用品", "高香", "4 元 × 6 把", "24.00"),
    ("开光用品", "酥油大拼蜡", "2 元 × 10 个", "20.00"),
    ("开光用品", "元宝", "5 元 × 9 袋", "45.00"),
    ("开光用品", "毛笔", "2 元 × 6 支", "12.00"),
    ("开光用品", "红布、黄布", "3 元 × 30 尺", "90.00"),
    ("开光仪式", "开光供品", "香蕉、桔子、苹果、花生、瓜子、桂圆、枣、蒸馍等", "319.00"),
    ("开光仪式", "开光费用（五尊像）及大庙原神像安放", "", "10,081.00"),
    ("开光仪式", "开光礼炮、鞭炮", "", "277.00"),
    ("庙内设施", "大庙安灯走线", "连工带料", "6,582.00"),
    ("庙内设施", "皮座椅", "8 个", "800.00"),
    ("庙内设施", "小方桌", "8 个", "200.00"),
    ("庙内设施", "大门铁锁", "8 元 × 2 把", "16.00"),
    ("庙内设施", "塑神仙用草帘", "3 卷", "100.00"),
    ("庙内设施", "优质水泥", "", "80.00"),
    ("石碑与木刻", "石碑（禅身）", "", "1,000.00"),
    ("石碑与木刻", "石碑（禅坐）", "", "400.00"),
    ("石碑与木刻", "石碑刻字", "900 × 2", "1,800.00"),
    ("石碑与木刻", "木雕对联、牌匾", "", "3,200.00"),
    ("石碑与木刻", "货拉拉运费", "木刻对联", "62.48"),
    ("资金经手", "冯彦庆处汇总", "安堂手 30,000；捐 10,000；飞飞手 50,000", "90,000.00"),
    ("庙内设施", "红绸", "2 元 × 30 米", "60.00"),
    ("庙内设施", "大庙窗户安玻璃", "工、料", "1,000.00"),
    ("庙内设施", "神像前供桌", "300 × 5 张", "1,500.00"),
    ("香炉等", "大香炉", "", "4,500.00"),
    ("香炉等", "小香炉", "1,600 × 2 个", "3,200.00"),
    ("香炉等", "香炉接站费", "", "300.00"),
    ("运输与其他", "铲车卸车等", "石碑 400；香炉 400；帮卸人员水、烟、啤酒 200", "1,000.00"),
    ("绿化与纪念品", "庙内种柏树", "3 棵", "320.00"),
    ("绿化与纪念品", "礼品水杯", "9.9 × 150 个", "1,485.00"),
    ("绿化与纪念品", "吉祥黄围巾", "3.5 × 200 条", "700.00"),
    ("绿化与纪念品", "包装红袋", "0.5 × 150", "75.00"),
    ("唱戏活动", "唱戏搭台等", "烧香、吉利烟酒红布、香、献品等", "300.00"),
    ("唱戏活动", "唱戏租凳子", "2 × 3 天 × 180 个", "1,080.00"),
    ("唱戏活动", "做饭棚", "", "180.00"),
    ("唱戏活动", "做饭用气", "150 × 4 罐", "600.00"),
    ("唱戏活动", "唱戏放礼花", "42 × 8 个（按所载金额）", "350.00"),
    ("唱戏活动", "唱戏用鞭炮", "45 × 10 个", "450.00"),
    ("唱戏活动", "防晒网", "", "160.00"),
    ("唱戏活动", "唱戏电费、条幅、对联等", "", "870.00"),
    ("庙内设施", "拜垫", "50 × 5 个（实收 250 元）", "250.00"),
]

TOTAL_INCOME = "217,900.00"
TOTAL_EXPENSE_AUDITED = "197,899.98"
REMAIN = "20,000.02"

# 大写金额（与上表审定数字一致）
_CN_INCOME = "贰拾壹万柒仟玖佰元整"
_CN_EXPENSE = "壹拾玖万柒仟捌佰玖拾玖元玖角捌分"
_CN_NET = "贰万元零贰分"

PUBLIC_SUMMARY_TEXT = (
    f"现将大庙收支数额公示如下：总收 {TOTAL_INCOME} 元（{_CN_INCOME}），"
    f"总支 {TOTAL_EXPENSE_AUDITED} 元（{_CN_EXPENSE}），"
    f"净余 {REMAIN} 元（{_CN_NET}），移交村委。"
)

SIGNED_NAMES = ("张鹏飞", "冯学兵", "梁小计", "王平文", "王志兴", "冯安堂")


def _parse_money(s: str) -> float:
    return float(s.replace(",", "").strip())


def _fmt_money(x: float) -> str:
    return f"{x:,.2f}"


def _html_esc(s: str) -> str:
    return s.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")


def merged_expense_table_html() -> str:
    """支出明细 HTML：类别列与小计列连续同类 rowspan 合并；最后一列为类别小计。"""
    lines = [
        "<table>",
        (
            "<thead><tr><th>类别</th><th>项目</th><th>明细/说明</th>"
            '<th style="text-align:right">金额（元）</th>'
            '<th style="text-align:right">类别小计（元）</th></tr></thead>'
        ),
        "<tbody>",
    ]
    i = 0
    n = len(ROWS)
    while i < n:
        cat = ROWS[i][0]
        j = i
        while j < n and ROWS[j][0] == cat:
            j += 1
        span = j - i
        block_sum = _fmt_money(sum(_parse_money(ROWS[k][3]) for k in range(i, j)))
        for k in range(i, j):
            _c, name, det, amt = ROWS[k]
            det = det if det else "—"
            nm, dt, am = _html_esc(name), _html_esc(det), _html_esc(amt)
            if k == i:
                if span > 1:
                    cat_td = (
                        f'<td rowspan="{span}" style="vertical-align:middle;white-space:nowrap;">'
                        f"{_html_esc(cat)}</td>"
                    )
                    sub_td = (
                        f'<td rowspan="{span}" style="vertical-align:middle;text-align:right">'
                        f"<strong>{block_sum}</strong></td>"
                    )
                else:
                    cat_td = (
                        '<td style="vertical-align:middle;white-space:nowrap;">'
                        f"{_html_esc(cat)}</td>"
                    )
                    sub_td = (
                        '<td style="vertical-align:middle;text-align:right">'
                        f"<strong>{block_sum}</strong></td>"
                    )
                lines.append(
                    f"<tr>{cat_td}"
                    f"<td>{nm}</td><td>{dt}</td>"
                    f'<td style="text-align:right">{am}</td>{sub_td}</tr>'
                )
            else:
                lines.append(
                    f"<tr><td>{nm}</td><td>{dt}</td>"
                    f'<td style="text-align:right">{am}</td></tr>'
                )
        i = j
    lines.append(
        '<tr><td colspan="4"><strong>合计（审定总支出）</strong></td>'
        f'<td style="text-align:right"><strong>{TOTAL_EXPENSE_AUDITED}</strong></td></tr>'
    )
    lines.append("</tbody></table>")
    return "\n".join(lines)


def _style_signature_table_no_lines(table) -> None:
    """签字表：去掉 Table Grid 样式残留，表级+单元格边框均为 none，并清掉底纹（避免 Word 仍画线/灰底）。"""
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    if tbl_pr is None:
        tbl_pr = OxmlElement("w:tblPr")
        tbl.insert(0, tbl_pr)
    ts = tbl_pr.find(qn("w:tblStyle"))
    if ts is not None:
        tbl_pr.remove(ts)
    old_tb = tbl_pr.find(qn("w:tblBorders"))
    if old_tb is not None:
        tbl_pr.remove(old_tb)
    tbl_borders = OxmlElement("w:tblBorders")
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"), "none")
        el.set(qn("w:sz"), "0")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "auto")
        tbl_borders.append(el)
    tbl_pr.append(tbl_borders)

    for row in table.rows:
        for cell in row.cells:
            tc_pr = cell._tc.get_or_add_tcPr()
            shd = tc_pr.find(qn("w:shd"))
            if shd is not None:
                tc_pr.remove(shd)
            old_cell_b = tc_pr.find(qn("w:tcBorders"))
            if old_cell_b is not None:
                tc_pr.remove(old_cell_b)
            tcb = OxmlElement("w:tcBorders")
            for side in ("top", "left", "bottom", "right"):
                el = OxmlElement(f"w:{side}")
                el.set(qn("w:val"), "none")
                el.set(qn("w:sz"), "0")
                el.set(qn("w:space"), "0")
                el.set(qn("w:color"), "auto")
                tcb.append(el)
            tc_pr.append(tcb)


def _fill_signature_grid_3x6(table, names: tuple[str, ...], *, n_rows: int = 3) -> None:
    """3 行 × 6 列：每列一人，每格一字，自上而下读为姓名。"""
    for col, name in enumerate(names):
        chars = list(name.strip())
        if len(chars) < n_rows:
            chars = [""] * (n_rows - len(chars)) + chars
        elif len(chars) > n_rows:
            chars = chars[:n_rows]
        for row in range(n_rows):
            cell = table.rows[row].cells[col]
            cell.text = chars[row]
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER


def _merge_consecutive_category_blocks(table, n_data_rows: int, col_idx: int) -> None:
    """将第 col_idx 列按「连续相同类别」块纵向合并（块内各行该列文本须已一致）。"""
    i = 0
    while i < n_data_rows:
        cat = ROWS[i][0]
        j = i + 1
        while j < n_data_rows and ROWS[j][0] == cat:
            j += 1
        row_top = i + 1
        row_bot = j
        if j - i > 1:
            table.cell(row_top, col_idx).merge(table.cell(row_bot, col_idx))
        i = j


def _set_run_cjk_font(run, font_name: str = CJK_FONT) -> None:
    """为中英文统一指定字体，并写入 eastAsia，避免中文显示为方框。"""
    run.font.name = font_name
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.get_or_add_rFonts()
    r_fonts.set(qn("w:ascii"), font_name)
    r_fonts.set(qn("w:hAnsi"), font_name)
    r_fonts.set(qn("w:eastAsia"), font_name)
    r_fonts.set(qn("w:cs"), font_name)


def _fix_paragraph_runs_font(paragraph, font_name: str = CJK_FONT) -> None:
    for run in paragraph.runs:
        _set_run_cjk_font(run, font_name)


def _iter_all_paragraphs(document: Document):
    for p in document.paragraphs:
        yield p
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    yield p


def fix_all_runs_cjk_font(document: Document, font_name: str = CJK_FONT) -> None:
    for p in _iter_all_paragraphs(document):
        _fix_paragraph_runs_font(p, font_name)


def _tighten_paragraph(paragraph) -> None:
    """单倍行距、压缩段前段后，减轻表格与正文「撑页」。"""
    pf = paragraph.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
    pf.line_spacing = 1.0
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)


def _tighten_all_paragraphs(document: Document) -> None:
    for p in document.paragraphs:
        _tighten_paragraph(p)
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    _tighten_paragraph(p)


def _header_row_bold(table) -> None:
    for cell in table.rows[0].cells:
        for p in cell.paragraphs:
            for r in p.runs:
                r.font.bold = True


def _right_align_money_column(table, col_index: int) -> None:
    for row in table.rows:
        cell = row.cells[col_index]
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT


def build() -> None:
    doc = Document()
    doc.styles["Normal"].font.size = Pt(11)

    t = doc.add_paragraph()
    tr = t.add_run("大庙捐款支持公示")
    tr.bold = True
    tr.font.size = Pt(18)
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_heading("说明", level=1)
    for text in [
        "1. 曾在群内公布，此处再公布一次，便于保存与转发。",
        "2. 庙上工程为塑像与广场一体，所收款项除塑像及相关必要开支外，结余移交村委。本次捐资为村委委托志兴、安堂、金盛三人挎名办理；村干部不直接塑像，故开支与余款权属归村委。",
        "3. 唱戏费用由村委支付，无欠账。",
    ]:
        doc.add_paragraph(text)

    doc.add_heading("收支总览（审定）", level=1)
    tbl = doc.add_table(rows=4, cols=2)
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, (a, b) in enumerate(
        [
            ("项目", "金额（元）"),
            ("总收入", TOTAL_INCOME),
            ("总支出", TOTAL_EXPENSE_AUDITED),
            ("结余（移交村委）", REMAIN),
        ]
    ):
        tbl.rows[i].cells[0].text = a
        tbl.rows[i].cells[1].text = b
    _header_row_bold(tbl)
    _right_align_money_column(tbl, 1)

    doc.add_heading("支出明细表", level=1)
    table = doc.add_table(rows=1 + len(ROWS) + 1, cols=5)
    table.style = "Table Grid"
    for j, h in enumerate(["类别", "项目", "明细/说明", "金额（元）", "类别小计（元）"]):
        table.rows[0].cells[j].text = h
    gi = 0
    while gi < len(ROWS):
        gj = gi + 1
        while gj < len(ROWS) and ROWS[gj][0] == ROWS[gi][0]:
            gj += 1
        block_sub = _fmt_money(sum(_parse_money(ROWS[k][3]) for k in range(gi, gj)))
        for idx in range(gi, gj):
            row = table.rows[idx + 1].cells
            cat, name, detail, amt = ROWS[idx]
            # 合并前只在块内首行写「类别」「小计」，否则 Word 纵向合并后会保留多段文字，看起来像一串重复数字
            row[0].text = cat if idx == gi else ""
            row[1].text = name
            row[2].text = detail if detail else "—"
            row[3].text = amt
            row[4].text = block_sub if idx == gi else ""
        gi = gj
    total_row_idx = len(ROWS) + 1
    _header_row_bold(table)
    _right_align_money_column(table, 3)
    _right_align_money_column(table, 4)
    _merge_consecutive_category_blocks(table, len(ROWS), 0)
    _merge_consecutive_category_blocks(table, len(ROWS), 4)
    table.cell(total_row_idx, 0).merge(table.cell(total_row_idx, 3))
    tot_cells = table.rows[total_row_idx].cells
    tot_cells[0].text = "合计（审定总支出）"
    tot_cells[1].text = TOTAL_EXPENSE_AUDITED

    doc.add_heading("数额与大写公示", level=1)
    doc.add_paragraph(PUBLIC_SUMMARY_TEXT)
    sig_tbl = doc.add_table(rows=3, cols=len(SIGNED_NAMES))
    sig_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    _fill_signature_grid_3x6(sig_tbl, SIGNED_NAMES)
    _style_signature_table_no_lines(sig_tbl)

    doc.add_heading("审议与日期", level=1)
    doc.add_paragraph("以上条款经当事 6 人审议（冯金盛因未打通电话除外），一致通过并签字。")
    doc.add_paragraph("日期：2025 年农历五月廿一日")

    for section in doc.sections:
        section.top_margin = Cm(1.7)
        section.bottom_margin = Cm(1.7)
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)

    fix_all_runs_cjk_font(doc, CJK_FONT)
    _tighten_all_paragraphs(doc)

    out = OUT_DIR / "大庙捐款支持公示.docx"
    doc.save(str(out))
    print("Wrote", out)


if __name__ == "__main__":
    import sys

    if len(sys.argv) > 1 and sys.argv[1] == "--print-expense-html":
        print(merged_expense_table_html())
    elif len(sys.argv) > 1 and sys.argv[1] == "--sync-md-table":
        md_path = OUT_DIR / "大庙捐款支持公示.md"
        md = md_path.read_text(encoding="utf-8")
        start = md.index("## 支出明细表")
        mark = "</tbody></table>"
        end = md.index(mark) + len(mark)
        tail = md[end:]
        md_path.write_text(
            md[:start] + "## 支出明细表\n\n" + merged_expense_table_html() + tail,
            encoding="utf-8",
        )
        print("Updated", md_path)
    else:
        build()
